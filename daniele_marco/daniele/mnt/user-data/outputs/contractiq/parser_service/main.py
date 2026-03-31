"""
contractiq/parser_service/main.py

Document Parser Service.
Reuses the team's existing parser logic (pdfplumber + python-docx).
Adds: model selector logic, page counting, cost estimation.
"""

import logging
import os
import re
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="ContractIQ Parser Service", version="1.0.0")

# Legal terms that trigger Sonnet instead of Haiku
COMPLEX_LEGAL_TERMS = [
    "penalty", "penale", "liability", "responsabilità", "termination", "risoluzione",
    "indemnification", "indennizzo", "warranty", "garanzia", "force majeure",
    "caso fortuito", "intellectual property", "proprietà intellettuale",
    "governing law", "legge applicabile", "arbitration", "arbitrato",
    "confidentiality", "riservatezza", "liquidated damages", "danni liquidati",
]

# Cost per token (Local Ollama)
INPUT_COST_PER_TOKEN = 0.0
OUTPUT_COST_PER_TOKEN = 0.0


class ParseRequest(BaseModel):
    file_path: str
    filename: str


@app.get("/health")
async def health():
    return {"status": "ok", "service": "parser_service"}


@app.post("/parse")
async def parse_document(req: ParseRequest):
    """Parse a document and return structured text with metadata."""
    file_path = Path(req.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {req.file_path}")

    ext = file_path.suffix.lower()

    if ext == ".pdf":
        text, pages = parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text, pages = parse_docx(file_path)
    elif ext == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        pages = max(1, len(text) // 3000)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    # Model selection logic (from team's model_selector.py)
    word_count = len(text.split())
    has_complex_terms = any(term.lower() in text.lower() for term in COMPLEX_LEGAL_TERMS)
    model_used = select_model(word_count, has_complex_terms)

    # Estimate cost (rough: 1 word ≈ 1.3 tokens)
    estimated_tokens = int(word_count * 1.3)
    estimated_cost = estimated_tokens * INPUT_COST_PER_TOKEN

    logger.info(f"Parsed {req.filename}: {pages} pages, {word_count} words, model={model_used}")

    return {
        "text": text,
        "pages": pages,
        "word_count": word_count,
        "model_used": model_used,
        "has_complex_terms": has_complex_terms,
        "estimated_cost_usd": round(estimated_cost, 6),
        "filename": req.filename,
    }


def parse_pdf(file_path: Path) -> tuple[str, int]:
    """Extract text from PDF using pdfplumber."""
    pages_text = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            pages_text.append(page_text)
    full_text = "\n\n".join(pages_text)
    # Clean up excessive whitespace
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)
    full_text = re.sub(r' {2,}', ' ', full_text)
    return full_text, len(pages_text)


def parse_docx(file_path: Path) -> tuple[str, int]:
    """Extract text from DOCX preserving structure."""
    doc = DocxDocument(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)
    # Estimate pages: ~3000 chars per page
    pages = max(1, len(full_text) // 3000)
    return full_text, pages


def select_model(word_count: int, has_complex_terms: bool) -> str:
    """
    All processing is handled by the external LLM (Qwen on Regolo AI).
    The parser itself only extracts text — LLM calls are made by dspy_agents.
    """
    return os.getenv("EXTERNAL_LLM_MODEL", "qwen3.5-122b")
